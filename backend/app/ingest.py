"""Evidence loading for Tracepoint.

Pure, synchronous, and free of any Cognee import - this module just turns the
files in ``data/`` into provenance-wrapped text blocks that the engine hands to
``cognee.add``. Ported verbatim from the v1 backup (only the DATA_DIR import
source changed). Two deliberate design choices live here:

1. Read content, wrap with a provenance header, add as TEXT (never a file path).
   Embedding a ``SOURCE:`` header inside every chunk lets answers cite their
   origin and lets ``purge()`` find the exact document to forget.

2. Turn CSV rows into natural-language sentences. Header-aware sentences make the
   same entity strings (a person, an IP, a filename) appear consistently, which
   is what lets Cognee resolve them into single shared graph nodes - the basis of
   the cross-source, multi-hop reasoning this product is built on.

The loader reads EVERY file in ``data/`` so an analyst can drop in their own logs
without touching code.
"""

from __future__ import annotations

import csv
import re
from pathlib import Path
from typing import Callable, NamedTuple

from .settings import DATA_DIR


class EvidenceDoc(NamedTuple):
    """One ingest-ready evidence document."""

    filename: str  # e.g. "vpn_logs.csv" - also the purge() target key
    source_type: str  # human label used in the provenance header
    reliability: str  # trust level; load-bearing for the purge demo
    text: str  # full provenance-wrapped text handed to cognee.add


# filename -> (human source label, reliability). The reliability line primes the
# graph so the anonymous tip is a visibly weak source and the phishing/SOC note
# is authoritative - this is what makes the before/after purge contrast land.
SOURCE_TYPES: dict[str, tuple[str, str]] = {
    "vpn_logs.csv": ("VPN connection log", "system-generated (high)"),
    "file_access_logs.csv": ("File access / DLP log", "system-generated (high)"),
    "email_logs.csv": ("Email gateway log", "system-generated (high)"),
    "badge_logs.csv": ("Physical badge-access log", "system-generated (high)"),
    "cctv_events.csv": ("CCTV person-detection log", "system-generated (high)"),
    "hr_notes.txt": ("HR case note", "internal record (medium)"),
    "phishing_evidence.txt": (
        "SOC threat-intel / phishing incident note",
        "security-verified (high)",
    ),
    "anonymous_tip.txt": (
        "anonymous whistleblower tip",
        "UNVERIFIED / uncorroborated (low)",
    ),
}
_DEFAULT_SOURCE = ("evidence document", "unclassified (medium)")

# For analyst-uploaded files (not in the curated demo map above), derive a sane
# human source label from the extension. Reliability is "uploaded evidence" so it
# reads as caller-provided rather than a system-verified feed.
_EXT_SOURCE: dict[str, tuple[str, str]] = {
    ".csv": ("Tabular log (CSV)", "uploaded evidence (medium)"),
    ".txt": ("Text document", "uploaded evidence (medium)"),
    ".pdf": ("PDF document", "uploaded evidence (medium)"),
    ".docx": ("Word document", "uploaded evidence (medium)"),
    ".doc": ("Word document", "uploaded evidence (medium)"),
    ".md": ("Text document", "uploaded evidence (medium)"),
}


def source_meta_for(filename: str) -> tuple[str, str]:
    """(source label, reliability) for a file: curated demo map first, then a
    per-extension default, then a generic fallback."""
    return (
        SOURCE_TYPES.get(filename)
        or _EXT_SOURCE.get(Path(filename).suffix.lower())
        or _DEFAULT_SOURCE
    )


# --------------------------------------------------------------------------- #
# CSV -> narrative sentence templates (header-aware). Each takes a DictReader
# row and returns one English sentence. Unknown CSVs fall back to a generic
# "field is value" rendering so drop-in files still work.
# --------------------------------------------------------------------------- #
def _vpn_row(r: dict) -> str:
    return (
        f"At {r['timestamp']}, user '{r['user']}' had a VPN login attempt from "
        f"IP {r['source_ip']} (location: {r['location']}); result: {r['status']}"
        f" (session {r['session_id']})."
    )


def _file_row(r: dict) -> str:
    return (
        f"At {r['timestamp']}, user '{r['user']}' performed a {r['action']} action "
        f"on file '{r['filename']}' on host {r['host']} (session {r['session_id']})."
    )


def _email_row(r: dict) -> str:
    return (
        f"At {r['timestamp']}, an email was sent from {r['sender']} to "
        f"{r['recipient']} with subject '{r['subject']}' and an attachment of "
        f"{r['attachment_size_mb']} MB."
    )


def _badge_row(r: dict) -> str:
    return (
        f"At {r['timestamp']}, employee {r['employee']} badged {r['direction']} "
        f"at the '{r['door']}' access point."
    )


def _cctv_row(r: dict) -> str:
    return (
        f"At {r['timestamp']}, CCTV camera {r['camera_id']} in zone "
        f"'{r['zone']}' detected {r['detected_person']}."
    )


_CSV_TEMPLATES: dict[str, Callable[[dict], str]] = {
    "vpn_logs.csv": _vpn_row,
    "file_access_logs.csv": _file_row,
    "email_logs.csv": _email_row,
    "badge_logs.csv": _badge_row,
    "cctv_events.csv": _cctv_row,
}


# A column whose name or value looks like a timestamp, so an arbitrary uploaded
# CSV still produces "At <ts>, ..." lines the deterministic timeline can parse.
_TS_KEY = re.compile(r"(^|_)(timestamp|datetime|date|time|ts)$", re.IGNORECASE)
_TS_VALUE = re.compile(r"^\d{4}-\d{2}-\d{2}[t ]\d{2}:\d{2}(:\d{2})?", re.IGNORECASE)


def _find_timestamp(r: dict) -> tuple[str | None, str | None]:
    """Return (column, normalized ISO timestamp) if the row has a timestamp."""
    for key, value in r.items():
        v = str(value or "").strip()
        if not v:
            continue
        if _TS_KEY.search(key or "") or _TS_VALUE.match(v):
            # Normalize "YYYY-MM-DD HH:MM:SS" -> "YYYY-MM-DDTHH:MM:SS" so the
            # timeline's ISO regex matches uploaded feeds too.
            iso = v.replace(" ", "T", 1) if _TS_VALUE.match(v) else v
            return key, iso
    return None, None


def _generic_row(r: dict) -> str:
    ts_key, ts = _find_timestamp(r)
    rest = [
        f"{k.replace('_', ' ')} is {v}"
        for k, v in r.items()
        if v not in (None, "") and k != ts_key
    ]
    body = ", ".join(rest)
    if ts:
        # Lead with "At <ts>," so the deterministic timeline can pick it up.
        return f"At {ts}, {body}." if body else f"At {ts}."
    return "Record: " + body + "."


def csv_to_narrative(path: Path) -> str:
    """Render a CSV as one natural-language sentence per row (header-aware)."""
    template = _CSV_TEMPLATES.get(path.name, _generic_row)
    lines: list[str] = []
    with path.open(newline="", encoding="utf-8") as fh:
        for row in csv.DictReader(fh):
            try:
                lines.append(template(row))
            except KeyError:
                # Header didn't match the specific template - degrade gracefully.
                lines.append(_generic_row(row))
    return "\n".join(lines)


def _pdf_to_text(path: Path) -> str:
    """Extract text from a PDF, page by page (pypdf, pure-Python, offline)."""
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts: list[str] = []
    for i, page in enumerate(reader.pages, start=1):
        try:
            text = (page.extract_text() or "").strip()
        except Exception:
            text = ""
        if text:
            parts.append(f"[page {i}]\n{text}")
    return "\n\n".join(parts).strip() or "(no extractable text in this PDF)"


def _docx_to_text(path: Path) -> str:
    """Extract text from a .docx: paragraphs then table rows (python-docx)."""
    import docx  # python-docx

    document = docx.Document(str(path))
    parts: list[str] = [p.text.strip() for p in document.paragraphs if p.text and p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip() or "(no extractable text in this document)"


def _wrap(filename: str, source_type: str, reliability: str, body: str) -> str:
    """Attach the provenance header used for citation and purge-targeting."""
    return (
        f"SOURCE: {filename}\n"
        f"SOURCE TYPE: {source_type}\n"
        f"RELIABILITY: {reliability}\n"
        f"---\n"
        f"{body.strip()}\n"
    )


def read_evidence_file(path: Path) -> EvidenceDoc:
    """Read one file into a provenance-wrapped :class:`EvidenceDoc`.

    CSV rows become narrative sentences; PDF/DOCX are text-extracted in-process;
    everything else is read as UTF-8 text. In all cases the body is wrapped with
    the same SOURCE header so citations, timeline, and purge keep working.
    """
    source_type, reliability = source_meta_for(path.name)
    suffix = path.suffix.lower()
    if suffix == ".csv":
        body = csv_to_narrative(path)
    elif suffix == ".pdf":
        body = _pdf_to_text(path)
    elif suffix == ".docx":
        body = _docx_to_text(path)
    else:
        body = path.read_text(encoding="utf-8", errors="replace")
    return EvidenceDoc(
        filename=path.name,
        source_type=source_type,
        reliability=reliability,
        text=_wrap(path.name, source_type, reliability, body),
    )


def load_all_evidence(data_dir: Path | str = DATA_DIR) -> list[EvidenceDoc]:
    """Load EVERY file in ``data_dir`` (sorted, deterministic)."""
    data_dir = Path(data_dir)
    if not data_dir.is_dir():
        raise FileNotFoundError(f"Evidence directory not found: {data_dir}")

    files = sorted(p for p in data_dir.iterdir() if p.is_file() and not p.name.startswith("."))
    if not files:
        raise FileNotFoundError(f"No evidence files found in {data_dir}")

    return [read_evidence_file(p) for p in files]


if __name__ == "__main__":
    # Quick manual check: `python -m app.ingest`
    docs = load_all_evidence()
    print(f"Loaded {len(docs)} evidence files:\n")
    for d in docs:
        print(f"  - {d.filename:24s} [{d.reliability}]")
    print("\n--- preview: vpn_logs.csv ---")
    print(next(d.text for d in docs if d.filename == "vpn_logs.csv")[:600])
